from __future__ import annotations

import asyncio
import base64
import html
import io
import json
import os
import re
import time
import unicodedata
import uuid
import urllib.request
import zipfile
from collections import defaultdict
from enum import Enum
from pathlib import Path
from typing import Annotated, AsyncGenerator, Literal
from urllib.parse import quote

import mammoth
from bs4 import BeautifulSoup
from bs4.element import NavigableString, Tag
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from fastapi import FastAPI, File, HTTPException, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse
from openai import OpenAI
from pydantic import BaseModel, Field
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas

app = FastAPI(title="Word to Tiptap Converter")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

STYLE_MAP = """
paragraph[style-name='Title'] => h1:fresh
paragraph[style-name='Heading 1'] => h1:fresh
paragraph[style-name='Heading 2'] => h2:fresh
paragraph[style-name='Heading 3'] => h3:fresh
paragraph[style-name='Heading 4'] => h4:fresh
paragraph[style-name='Heading 5'] => h5:fresh
paragraph[style-name='Heading 6'] => h6:fresh
r[style-name='Strong'] => strong
r[style-name='Emphasis'] => em
"""

SUPPORTED_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/octet-stream",  # some browsers fallback
}

WORD_FILE_MIME_TYPES = SUPPORTED_MIME_TYPES | {"application/msword"}

BASE_DIR = Path(__file__).resolve().parent
ONLYOFFICE_STORAGE_DIR = BASE_DIR / "storage" / "onlyoffice"
ONLYOFFICE_STORAGE_DIR.mkdir(parents=True, exist_ok=True)

DOCUMENT_SERVER_INTERNAL = os.getenv(
    "ONLYOFFICE_DOCUMENT_SERVER_URL", "http://localhost:8085"
)
DOCUMENT_SERVER_PUBLIC = os.getenv(
    "ONLYOFFICE_DOCSERVER_PUBLIC_URL", DOCUMENT_SERVER_INTERNAL
)
ONLYOFFICE_BACKEND_BASE = os.getenv("ONLYOFFICE_PUBLIC_BASE_URL")
SUPPORTED_ONLYOFFICE_EXTENSIONS = {"docx", "doc", "pptx", "xlsx"}


class ConversionNote(BaseModel):
    type: str
    message: str


class ConversionResponse(BaseModel):
    html: str
    notes: list[ConversionNote] = []


class DiffStats(BaseModel):
    inserted_tokens: int
    deleted_tokens: int
    replaced_tokens: int


class DiffLocation(BaseModel):
    section_title: str | None = None
    block_summary: str | None = None


class DiffItem(BaseModel):
    id: str
    type: Literal["insert", "delete", "replace"]
    original_text: str
    modified_text: str
    original_location: DiffLocation | None = None
    modified_location: DiffLocation | None = None


class DiffResponse(BaseModel):
    original_html: str
    modified_html: str
    diff_html: str
    stats: DiffStats
    diff_items: list[DiffItem] = []
    original_notes: list[ConversionNote] = []
    modified_notes: list[ConversionNote] = []


class ExportRequest(BaseModel):
    content: str
    format: Literal["docx", "pdf", "json"]
    filename: str | None = None


class SensitiveHit(BaseModel):
    category: str
    field: str
    value: str
    count: int


class DesensitizeResponse(BaseModel):
    sanitized_docx: str
    filename: str
    total_hits: int
    hits: list[SensitiveHit] = []
    sanitized_preview: str | None = None


DIFF_TYPE_LABELS: dict[str, str] = {
    "insert": "新增",
    "delete": "删除",
    "replace": "修改",
}


class AiTransformRequest(BaseModel):
    mode: Literal["rewrite", "expand", "rephrase", "custom"]
    markdown: str
    user_instruction: str | None = None


class MarkdownPayload(BaseModel):
    markdown: str


class AiAction(str, Enum):
    GENERATE = "generate"
    REWRITE = "rewrite"
    EXPAND = "expand"
    CUSTOM = "custom"


RISK_DIMENSION_DEFINITIONS: dict[str, str] = {
    "DUTY_OBLIGATION_MATCH": "权利义务匹配，关注上下游责任、义务是否一致及可传递。",
    "LIQUIDATED_DAMAGES": "违约金条款，包括违约情形、计算方式、上限等。",
    "SETTLEMENT_CYCLE": "结算周期与付款条件，关注账期差异和回款风险。",
    "LIABILITY_ASSUMPTION": "责任承担主体及范围，是否存在我方单方兜底。",
    "LIABILITY_LIMITATION": "责任限制与赔偿上限，检查是否倒挂或缺失。",
    "EXEMPTION_CLAUSE": "免责条款，是否可对等适用于上下游。",
    "INSURANCE_BENEFICIARY": "保险受益人设置是否与风险主体一致。",
    "INSURANCE_AMOUNT_DEDUCTIBLE": "保险保额与免赔额是否足以覆盖下游承诺。",
    "WARRANTY_PERIOD": "质量保修期及责任期限。",
    "RETENTION_REFUND": "质保金留存与返还条件。",
    "PRICE_ADJUSTMENT": "价格调整机制及触发条件。",
    "SERVICE_EVALUATION": "服务评价与考核指标，关联违约或费用扣减。",
}


class RiskClause(BaseModel):
    id: str | None = None
    text: str
    heading_path: list[str] | None = None


class RiskDimensionPayload(BaseModel):
    dimension_code: str
    upstream_clauses: list[RiskClause] = Field(default_factory=list)
    downstream_clauses: list[RiskClause] = Field(default_factory=list)
    context_summary: str | None = None


class RiskTransferRequest(BaseModel):
    dimensions: list[RiskDimensionPayload]
    model: str = "gpt-4o-mini"
    temperature: float = 0.2
    language: Literal["zh", "en"] = "zh"


class RiskTransferItem(BaseModel):
    dimension_code: str
    severity: Literal["high", "medium", "low"]
    risk_transfer_status: Literal[
        "fully_transferred", "partially_transferred", "not_transferred", "inverted"
    ]
    risk_type: str
    explanation: str
    suggestion: str
    matched_upstream_clause_ids: list[str] = Field(default_factory=list)
    matched_downstream_clause_ids: list[str] = Field(default_factory=list)


class RiskTransferResponse(BaseModel):
    items: list[RiskTransferItem]


class RiskTransferAnalyzer:
    def __init__(
        self,
        *,
        client: OpenAI,
        model: str = "gpt-4o-mini",
        temperature: float = 0.2,
        language: str = "zh",
    ) -> None:
        self.client = client
        self.model = model
        self.temperature = temperature
        self.language = language

    async def analyze(self, payload: RiskTransferRequest) -> list[RiskTransferItem]:
        results: list[RiskTransferItem] = []
        for dimension in payload.dimensions:
            results.append(await self._analyze_dimension(dimension))
        return results

    async def _analyze_dimension(self, dimension: RiskDimensionPayload) -> RiskTransferItem:
        messages = self._build_messages(dimension)

        def _call_openai() -> str:
            completion = self.client.chat.completions.create(
                model=self.model,
                temperature=self.temperature,
                response_format={"type": "json_object"},
                messages=messages,
            )
            return completion.choices[0].message.content or "{}"

        try:
            content = await asyncio.to_thread(_call_openai)
        except Exception as exc:  # pragma: no cover - depends on network credentials
            raise HTTPException(
                status_code=502,
                detail="Failed to call OpenAI for risk transfer analysis",
            ) from exc

        try:
            payload = json.loads(content)
        except json.JSONDecodeError as exc:  # pragma: no cover - defensive
            raise HTTPException(
                status_code=500,
                detail="Unable to parse OpenAI response for risk transfer analysis",
            ) from exc

        return RiskTransferItem(
            dimension_code=dimension.dimension_code,
            severity=payload.get("severity", "medium"),
            risk_transfer_status=payload.get(
                "risk_transfer_status", "partially_transferred"
            ),
            risk_type=payload.get("risk_type", "unspecified"),
            explanation=payload.get("explanation", ""),
            suggestion=payload.get("suggestion", ""),
            matched_upstream_clause_ids=payload.get("matched_upstream_clause_ids", []),
            matched_downstream_clause_ids=payload.get("matched_downstream_clause_ids", []),
        )

    def _build_messages(self, dimension: RiskDimensionPayload) -> list[dict[str, object]]:
        dimension_label = RISK_DIMENSION_DEFINITIONS.get(
            dimension.dimension_code, "其他风险维度"
        )
        language_hint = "请使用中文输出。" if self.language == "zh" else "Please respond in English."
        system_prompt = f"""
你是资深的上下游合同风险审核专家，专注判断风险能否顺利从下游向上游转嫁。
输出要求：
- 仅输出 JSON，不要添加多余文字。
- severity 取值：high / medium / low。
- risk_transfer_status 取值：fully_transferred / partially_transferred / not_transferred / inverted。
- 解释中要直接指出上下游差异、倒挂或缺失点，并给出可执行的修改建议。

十二个固定维度编码及释义：
{json.dumps(RISK_DIMENSION_DEFINITIONS, ensure_ascii=False, indent=2)}
{language_hint}
""".strip()

        def _format_clause(clause: RiskClause, index: int) -> dict[str, str | list[str] | None]:
            return {
                "id": clause.id or f"c-{index}",
                "text": clause.text,
                "heading_path": clause.heading_path,
            }

        upstream = [_format_clause(clause, idx) for idx, clause in enumerate(dimension.upstream_clauses, 1)]
        downstream = [
            _format_clause(clause, idx) for idx, clause in enumerate(dimension.downstream_clauses, 1)
        ]

        user_payload = {
            "task": "risk_transfer_assessment",
            "dimension_code": dimension.dimension_code,
            "dimension_label": dimension_label,
            "context_summary": dimension.context_summary,
            "upstream_clauses": upstream,
            "downstream_clauses": downstream,
            "expected_fields": {
                "severity": "high/medium/low",
                "risk_transfer_status": "fully_transferred/partially_transferred/not_transferred/inverted",
                "risk_type": "简短的风险标签，如 liability_cap_mismatch、missing_insurance 等",
                "explanation": "200 字以内，指出风险原因、责任链断点、条款缺失或倒挂。",
                "suggestion": "150 字以内，提供可执行的修改建议。",
                "matched_upstream_clause_ids": "引用上游条款 id 列表，用于前端高亮。",
                "matched_downstream_clause_ids": "引用下游条款 id 列表，用于前端高亮。",
            },
        }

        return [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": json.dumps(user_payload, ensure_ascii=False, indent=2)},
        ]


def _get_openai_client() -> OpenAI:
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise HTTPException(
            status_code=500,
            detail="OPENAI_API_KEY is required to run risk transfer analysis.",
        )

    base_url = os.getenv("OPENAI_BASE_URL")
    if base_url:
        return OpenAI(api_key=api_key, base_url=base_url)
    return OpenAI(api_key=api_key)


def _format_sse(*, data: str, event: str | None = None, event_id: str | None = None) -> str:
    """Format payload in SSE wire format."""

    lines: list[str] = []
    if event_id:
        lines.append(f"id: {event_id}")
    if event:
        lines.append(f"event: {event}")
    payload_lines = data.splitlines() or [""]
    lines.extend(f"data: {line}" for line in payload_lines)
    lines.append("")
    return "\n".join(lines)


def _chunk_text(text: str, chunk_size: int = 48) -> list[str]:
    """Split text into balanced chunks for streaming."""

    if chunk_size <= 0:
        return [text]
    return [text[i : i + chunk_size] for i in range(0, len(text), chunk_size)] or [""]


def _simulate_ai_response(action: AiAction, text: str, instruction: str = "") -> str:
    clean_text = text.strip()
    if not clean_text:
        return "请选择一段文本后再试。"

    if action is AiAction.GENERATE:
        return (
            f"基于「{clean_text}」生成的新句子：为确保条款清晰，"
            "双方应在签署后十个工作日内完成约定事项。"
        )

    if action is AiAction.REWRITE:
        rewritten = (
            clean_text.replace("应当", "应")
            .replace("不得", "严禁")
            .replace("立即", "立刻")
            .replace("双方", "双方各方")
            .replace("保证", "确保")
        )
        if rewritten == clean_text:
            return f"经优化表述：{clean_text}"
        return rewritten

    if action is AiAction.EXPAND:
        return (
            f"{clean_text}。为提升条款的可执行性，建议补充具体时间节点、责任划分"
            "以及必要的沟通机制，确保各项义务能够有效落实。"
        )

    if action is AiAction.CUSTOM:
        safe_instruction = instruction.strip() or "自定义指令"
        return (
            f"根据「{safe_instruction}」调整后的内容：{clean_text}，"
            "在保持原意的基础上补充格式和细节，确保条款表述清晰可执行。"
        )

    return clean_text


@app.post("/api/risk-transfer/analyze", response_model=RiskTransferResponse)
async def analyze_risk_transfer(request: RiskTransferRequest) -> RiskTransferResponse:
    """Call OpenAI to judge whether risks are transferred across upstream/downstream contracts."""

    client = _get_openai_client()
    analyzer = RiskTransferAnalyzer(
        client=client,
        model=request.model,
        temperature=request.temperature,
        language=request.language,
    )

    items = await analyzer.analyze(request)
    return RiskTransferResponse(items=items)


@app.get("/ai/editor/stream")
async def stream_ai_editor(
    action: AiAction,
    text: str = "",
    instruction: str = "",
    request_id: str | None = None,
) -> StreamingResponse:
    """Stream AI results for the editor via SSE."""

    clean_text = text.strip()
    resolved_request_id = request_id or uuid.uuid4().hex

    async def event_publisher() -> AsyncGenerator[str, None]:
        yield "retry: 3000\n\n"
        start_payload = {
            "requestId": resolved_request_id,
            "action": action.value,
            "receivedText": clean_text,
            "instruction": instruction,
        }
        yield _format_sse(
            data=json.dumps(start_payload, ensure_ascii=False),
            event="start",
            event_id=f"{resolved_request_id}:0",
        )

        if not clean_text:
            done_payload = {
                "requestId": resolved_request_id,
                "status": "empty",
                "message": "请选择一段文本后再试。",
            }
            yield _format_sse(
                data=json.dumps(done_payload, ensure_ascii=False),
                event="done",
                event_id=f"{resolved_request_id}:1",
            )
            return

        ai_result = _simulate_ai_response(action, clean_text, instruction)
        chunks = _chunk_text(ai_result)

        for index, chunk in enumerate(chunks, start=1):
            chunk_payload = {
                "requestId": resolved_request_id,
                "content": chunk,
                "index": index,
                "total": len(chunks),
            }
            yield _format_sse(
                data=json.dumps(chunk_payload, ensure_ascii=False),
                event="chunk",
                event_id=f"{resolved_request_id}:{index}",
            )
            await asyncio.sleep(0.18)

        done_payload = {
            "requestId": resolved_request_id,
            "status": "completed",
            "result": ai_result,
            "totalChunks": len(chunks),
            "meta": {"instruction": instruction},
        }
        yield _format_sse(
            data=json.dumps(done_payload, ensure_ascii=False),
            event="done",
            event_id=f"{resolved_request_id}:{len(chunks) + 1}",
        )

    headers = {
        "Cache-Control": "no-cache",
        "X-Accel-Buffering": "no",
    }
    return StreamingResponse(event_publisher(), media_type="text/event-stream", headers=headers)


def _ensure_docx(file: UploadFile) -> None:
    if not file.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Only .docx files are supported")

    if file.content_type not in SUPPORTED_MIME_TYPES:
        raise HTTPException(status_code=400, detail="Unsupported file type")


async def _convert_to_html(file: UploadFile) -> tuple[str, list[ConversionNote]]:
    _ensure_docx(file)

    raw_bytes = await file.read()
    if not raw_bytes:
        raise HTTPException(status_code=400, detail="Uploaded file is empty")

    try:
        with io.BytesIO(raw_bytes) as buffer:
            result = mammoth.convert_to_html(buffer, style_map=STYLE_MAP)
    except Exception as exc:  # pragma: no cover - defensive
        raise HTTPException(status_code=500, detail="Failed to process document") from exc

    html_content = result.value.strip()

    notes: list[ConversionNote] = []
    if result.messages:
        notes = [
            ConversionNote(type=message.type, message=message.message)
            for message in result.messages
        ]

    return html_content, notes


def _is_cjk_char(char: str) -> bool:
    if not char:
        return False
    codepoint = ord(char)
    return (
        0x4E00 <= codepoint <= 0x9FFF  # CJK Unified Ideographs
        or 0x3400 <= codepoint <= 0x4DBF  # CJK Unified Ideographs Extension A
        or 0x20000 <= codepoint <= 0x2A6DF  # CJK Unified Ideographs Extension B
        or 0x2A700 <= codepoint <= 0x2B73F  # CJK Unified Ideographs Extension C
        or 0x2B740 <= codepoint <= 0x2B81F  # CJK Unified Ideographs Extension D
        or 0x2B820 <= codepoint <= 0x2CEAF  # CJK Unified Ideographs Extension E
        or 0xF900 <= codepoint <= 0xFAFF  # CJK Compatibility Ideographs
    )


def _is_punctuation(char: str) -> bool:
    return unicodedata.category(char).startswith("P")


def _tokenize_text(text: str) -> list[str]:
    tokens: list[str] = []
    buffer: list[str] = []

    def flush_buffer() -> None:
        if buffer:
            tokens.append("".join(buffer))
            buffer.clear()

    for char in text:
        if char.isspace():
            flush_buffer()
            tokens.append(char)
            continue

        if _is_cjk_char(char) or _is_punctuation(char):
            flush_buffer()
            tokens.append(char)
            continue

        buffer.append(char)

    flush_buffer()

    return tokens


def _prepare_html_tokens(html_content: str) -> tuple[BeautifulSoup, list[str], list[dict[str, object]]]:
    soup = BeautifulSoup(html_content or "", "html.parser")
    tokens: list[str] = []
    node_infos: list[dict[str, object]] = []

    for node in soup.descendants:
        if not isinstance(node, NavigableString):
            continue

        text = str(node)
        if text == "":
            continue

        parts = _tokenize_text(text)
        if not parts:
            continue

        start_index = len(tokens)
        tokens.extend(parts)
        node_infos.append(
            {
                "node": node,
                "tokens": parts,
                "start": start_index,
                "end": len(tokens),
            }
        )

    return soup, tokens, node_infos


def _escape_tokens(tokens: list[str]) -> str:
    return "".join(html.escape(token) for token in tokens)


def _truncate_text(value: str, limit: int = 80) -> str:
    if len(value) <= limit:
        return value
    return value[: limit - 1] + "…"


def _mask_value(value: str) -> str:
    if not value:
        return ""

    masked = []
    for char in value:
        if char.isspace():
            masked.append(char)
        elif char.isdigit():
            masked.append("•")
        elif char.isalpha():
            masked.append("*")
        elif _is_cjk_char(char):
            masked.append("＊")
        else:
            masked.append("※")
    return "".join(masked) or "＊"


def _mask_text_with_map(text: str, mask_map: dict[str, str]) -> str:
    if not text or not mask_map:
        return text

    sanitized = text
    for raw, masked in sorted(mask_map.items(), key=lambda item: len(item[0]), reverse=True):
        sanitized = re.sub(re.escape(raw), masked, sanitized)
    return sanitized


def _collect_docx_text(raw_bytes: bytes) -> str:
    try:
        document = Document(io.BytesIO(raw_bytes))
    except Exception as exc:  # pragma: no cover - defensive
        raise HTTPException(status_code=400, detail="无法读取合同内容，请确认文件是否为有效的 Word 文档") from exc

    texts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text:
            texts.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    texts.append(cell.text)

    return "\n".join(texts)


def _find_sensitive_hits(text: str) -> list[SensitiveHit]:
    hit_map: dict[tuple[str, str], SensitiveHit] = {}

    for config in SENSITIVE_FIELD_CONFIGS:
        patterns: list[PatternGroup] = config.get("patterns", [])  # type: ignore[assignment]
        for pattern, group_index in patterns:
            for match in pattern.finditer(text):
                try:
                    value = match.group(group_index) if group_index else match.group(0)
                except IndexError:  # pragma: no cover - defensive
                    value = match.group(0)

                cleaned_value = value.strip()
                if not cleaned_value:
                    continue

                key = (str(config["field"]), cleaned_value)
                if key not in hit_map:
                    hit_map[key] = SensitiveHit(
                        category=str(config["category"]),
                        field=str(config["field"]),
                        value=cleaned_value,
                        count=1,
                    )
                else:
                    hit_map[key].count += 1

    return sorted(hit_map.values(), key=lambda hit: (hit.category, hit.field, hit.value))


def _sanitize_docx_bytes(docx_bytes: bytes, mask_map: dict[str, str]) -> bytes:
    if not mask_map:
        return docx_bytes

    input_buffer = io.BytesIO(docx_bytes)
    output_buffer = io.BytesIO()

    with zipfile.ZipFile(input_buffer, "r") as source_zip, zipfile.ZipFile(
        output_buffer, "w"
    ) as target_zip:
        for item in source_zip.infolist():
            data = source_zip.read(item.filename)
            if item.filename.endswith(".xml"):
                xml_text = data.decode("utf-8")
                masked_text = _mask_text_with_map(xml_text, mask_map)
                data = masked_text.encode("utf-8")
            target_zip.writestr(item, data)

    return output_buffer.getvalue()


async def _read_word_file_bytes(file: UploadFile) -> tuple[bytes, str]:
    filename = (file.filename or "合同.docx").strip()
    lowered = filename.lower()

    if not lowered.endswith((".docx", ".doc")):
        raise HTTPException(status_code=400, detail="仅支持上传 .doc 或 .docx 合同文件")

    if file.content_type not in WORD_FILE_MIME_TYPES:
        raise HTTPException(status_code=400, detail="不支持的文件类型")

    raw_bytes = await file.read()
    if not raw_bytes:
        raise HTTPException(status_code=400, detail="上传文件为空")

    if lowered.endswith(".doc"):
        raise HTTPException(status_code=400, detail="暂不支持 .doc，请先转为 .docx 后再试")

    return raw_bytes, filename


PatternGroup = tuple[re.Pattern[str], int]


SENSITIVE_FIELD_CONFIGS: list[dict[str, object]] = [
    {
        "category": "主体身份信息",
        "field": "企业名称",
        "patterns": [
            (
                re.compile(
                    r"[\u4e00-\u9fa5A-Za-z0-9（）()·]{2,}(?:有限责任公司|股份有限公司|有限公司|集团|公司|合伙企业|工作室|事务所)"
                ),
                0,
            ),
        ],
    },
    {
        "category": "人名",
        "field": "人名",
        "patterns": [(re.compile(r"(?<![\w])[\u4e00-\u9fa5]{2,4}(?:先生|女士)?"), 0)],
    },
    {
        "category": "主体身份信息",
        "field": "身份证号",
        "patterns": [
            (
                re.compile(
                    r"\b\d{6}(?:19|20)?\d{2}(?:0[1-9]|1[0-2])(?:0[1-9]|[12]\d|3[01])\d{3}[0-9Xx]\b"
                ),
                0,
            )
        ],
    },
    {
        "category": "联系方式",
        "field": "联系电话",
        "patterns": [
            (
                re.compile(r"\b1[3-9]\d{9}\b"),
                0,
            ),
            (
                re.compile(r"\b0\d{2,3}-?\d{7,8}\b"),
                0,
            ),
        ],
    },
    {
        "category": "联系方式",
        "field": "邮箱",
        "patterns": [(re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}"), 0)],
    },
    {
        "category": "地址",
        "field": "地址",
        "patterns": [
            (
                re.compile(
                    r"[\u4e00-\u9fa5A-Za-z0-9]{2,}(?:省|市|自治区|区|县|镇|乡|街道|大道|路|街|巷|号)[\u4e00-\u9fa5A-Za-z0-9#\-（）()]{2,40}"
                ),
                0,
            )
        ],
    },
    {
        "category": "企业注册信息",
        "field": "统一社会信用代码",
        "patterns": [(re.compile(r"\b[0-9A-Z]{18}\b"), 0)],
    },
    {
        "category": "企业注册信息",
        "field": "法定代表人",
        "patterns": [(re.compile(r"法定代表人[:：]?\s*([\u4e00-\u9fa5]{2,4})"), 1)],
    },
    {
        "category": "银行与税务信息",
        "field": "银行账户",
        "patterns": [(re.compile(r"\b\d{12,24}\b"), 0)],
    },
    {
        "category": "银行与税务信息",
        "field": "开户银行",
        "patterns": [
            (
                re.compile(r"[\u4e00-\u9fa5A-Za-z]{2,}(?:银行|信用社|合作社)[\u4e00-\u9fa5A-Za-z]*"),
                0,
            )
        ],
    },
    {
        "category": "银行与税务信息",
        "field": "纳税人识别号",
        "patterns": [(re.compile(r"\b[0-9A-Z]{15,20}\b"), 0)],
    },
    {
        "category": "合同与项目标识",
        "field": "合同编号",
        "patterns": [(re.compile(r"合同编号[:：]?\s*([A-Za-z0-9\-]{4,})"), 1)],
    },
    {
        "category": "合同与项目标识",
        "field": "项目名称",
        "patterns": [
            (
                re.compile(r"项目名称[:：]?\s*([\u4e00-\u9fa5A-Za-z0-9（）()·\-]{2,})"),
                1,
            )
        ],
    },
    {
        "category": "时间信息",
        "field": "日期",
        "patterns": [
            (
                re.compile(r"\b\d{4}年\d{1,2}月\d{1,2}日\b"),
                0,
            ),
            (
                re.compile(r"\b\d{4}-\d{1,2}-\d{1,2}\b"),
                0,
            ),
        ],
    },
    {
        "category": "时间段",
        "field": "时间段",
        "patterns": [
            (
                re.compile(r"\b\d{4}年\d{1,2}月\d{1,2}日\s*[至\-]+\s*\d{4}年\d{1,2}月\d{1,2}日\b"),
                0,
            )
        ],
    },
]


EXPORT_BLOCK_TAGS: tuple[str, ...] = (
    "h1",
    "h2",
    "h3",
    "h4",
    "h5",
    "h6",
    "p",
    "li",
    "blockquote",
    "pre",
    "code",
    "table",
)


def _sanitize_filename(value: str | None) -> str:
    if not value:
        return "合同导入编辑"
    sanitized = re.sub(r"[\\/:*?\"<>|]", "_", value)
    sanitized = re.sub(r"\s+", " ", sanitized).strip()
    return sanitized or "合同导入编辑"


def _html_to_plaintext_lines(html_content: str) -> list[str]:
    soup = BeautifulSoup(html_content or "", "html.parser")
    body = soup.body or soup
    lines: list[str] = []

    for block in body.find_all(EXPORT_BLOCK_TAGS):
        separator = "\n" if block.name in {"pre", "code"} else " "
        text = block.get_text(separator, strip=True)
        if text == "":
            lines.append("")
            continue
        parts = text.splitlines() or [""]
        lines.extend(parts)

    if not lines:
        fallback = (body.get_text("\n", strip=True) or "").splitlines()
        if fallback:
            lines.extend(fallback)
        else:
            lines.append("")

    return [line.rstrip("\r") for line in lines]


def _render_docx_document(html_content: str) -> io.BytesIO:
    document = Document()
    soup = BeautifulSoup(html_content or "", "html.parser")
    body = soup.body or soup
    blocks = body.find_all(EXPORT_BLOCK_TAGS)

    if not blocks:
        document.add_paragraph("")

    for block in blocks:
        separator = "\n" if block.name in {"pre", "code"} else " "
        text = block.get_text(separator, strip=True)

        if block.name.startswith("h") and len(block.name) == 2 and block.name[1].isdigit():
            level = max(0, min(int(block.name[1]) - 1, 4))
            document.add_heading(text or "", level=level)
            continue

        if block.name == "li":
            parent = block.parent if isinstance(block.parent, Tag) else None
            style = "List Number" if parent and parent.name == "ol" else "List Bullet"
            document.add_paragraph(text or "", style=style)
            continue

        if block.name in {"pre", "code"}:
            paragraph = document.add_paragraph()
            content = text.splitlines() or [""]
            for index, line in enumerate(content):
                run = paragraph.add_run(line)
                run.font.name = "Courier New"
                run.font.size = Pt(10)
                if index < len(content) - 1:
                    run.add_break()
            if not content:
                paragraph.add_run("")
            continue

        if block.name == "blockquote":
            paragraph = document.add_paragraph(text or "")
            if "Intense Quote" in document.styles:
                paragraph.style = "Intense Quote"
            continue

        if block.name == "table":
            rows = block.find_all("tr")
            if not rows:
                continue
            max_cols = max((len(row.find_all(["th", "td"])) for row in rows), default=0)
            if max_cols == 0:
                continue
            table = document.add_table(rows=len(rows), cols=max_cols)
            for row_index, row in enumerate(rows):
                cells = row.find_all(["th", "td"])
                for col_index, cell in enumerate(cells):
                    if col_index >= max_cols:
                        break
                    table.cell(row_index, col_index).text = cell.get_text(" ", strip=True)
            document.add_paragraph("")
            continue

        document.add_paragraph(text or "")

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def _render_pdf_document(html_content: str) -> io.BytesIO:
    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    _, height = A4
    margin = 25 * mm
    text_object = pdf.beginText(margin, height - margin)
    text_object.setFont("Helvetica", 12)

    for line in _html_to_plaintext_lines(html_content):
        if text_object.getY() <= margin:
            pdf.drawText(text_object)
            pdf.showPage()
            text_object = pdf.beginText(margin, height - margin)
            text_object.setFont("Helvetica", 12)
        text_object.textLine(line)

    pdf.drawText(text_object)
    pdf.save()
    buffer.seek(0)
    return buffer


def _render_json_document(html_content: str) -> io.BytesIO:
    lines = _html_to_plaintext_lines(html_content)
    plain_text = "\n".join(lines)
    payload = {
        "html": html_content or "",
        "plain_text": plain_text,
        "character_count": len(plain_text.replace("\n", "")),
        "line_count": len(lines),
    }
    buffer = io.BytesIO(json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8"))
    buffer.seek(0)
    return buffer


def _build_file_response(data: bytes, media_type: str, filename: str) -> StreamingResponse:
    response = StreamingResponse(iter([data]), media_type=media_type)
    response.headers["Content-Length"] = str(len(data))
    utf8_filename = quote(filename)
    response.headers[
        "Content-Disposition"
    ] = f"attachment; filename=\"{filename}\"; filename*=UTF-8''{utf8_filename}"
    return response


def _ensure_onlyoffice_file(file_id: str) -> Path:
    file_path = ONLYOFFICE_STORAGE_DIR / file_id
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="文件不存在或已被删除。")
    return file_path


def _resolve_backend_base(request: Request) -> str:
    if ONLYOFFICE_BACKEND_BASE:
        return ONLYOFFICE_BACKEND_BASE.rstrip("/")
    return str(request.base_url).rstrip("/")


async def _persist_onlyoffice_upload(file: UploadFile) -> tuple[str, str]:
    filename = (file.filename or "").strip()
    if not filename:
        raise HTTPException(status_code=400, detail="请上传有效的文件。")

    extension = Path(filename).suffix.lower().lstrip(".")
    if extension not in SUPPORTED_ONLYOFFICE_EXTENSIONS:
        raise HTTPException(status_code=400, detail="暂不支持该文件类型，请上传 Office 文档。")

    safe_name = re.sub(r"[^A-Za-z0-9._-]", "_", Path(filename).name) or f"document.{extension}"
    file_id = f"{uuid.uuid4().hex}_{safe_name}"

    contents = await file.read()
    if not contents:
        raise HTTPException(status_code=400, detail="文件内容为空，请重新上传。")

    target_path = ONLYOFFICE_STORAGE_DIR / file_id
    target_path.write_bytes(contents)
    return file_id, safe_name


def _build_onlyoffice_config(file_id: str, display_name: str, request: Request) -> dict[str, object]:
    file_path = _ensure_onlyoffice_file(file_id)
    file_type = Path(display_name).suffix.lstrip(".") or "docx"

    if file_type in {"xlsx"}:
        document_type = "spreadsheet"
    elif file_type in {"pptx"}:
        document_type = "presentation"
    else:
        document_type = "word"

    base_url = _resolve_backend_base(request)
    file_url = f"{base_url}/onlyoffice/files/{file_id}"
    callback_url = f"{base_url}/onlyoffice/callback/{file_id}"

    last_modified = int(file_path.stat().st_mtime) if file_path.exists() else int(time.time())
    return {
        "document": {
            "fileType": file_type,
            "title": display_name,
            "key": f"{file_id}-{last_modified}",
            "url": file_url,
            "permissions": {
                "comment": False,
                "download": True,
                "edit": True,
                "print": True,
                "review": False,
            },
        },
        "documentType": document_type,
        "editorConfig": {
            "mode": "edit",
            "callbackUrl": callback_url,
            "lang": "zh-CN",
            "user": {
                "id": "contract-user",
                "name": "合同协作用户",
            },
            "customization": {
                "autosave": True,
                "compactHeader": True,
                "compactToolbar": True,
                "hideRightMenu": True,
                "leftMenu": False,
                "rightMenu": False,
                "toolbarHideFileName": True,
                "toolbarNoTabs": True,
                "feedback": False,
                "help": False,
                "chat": False,
                "comments": False,
                "zoom": 100,
                "showReviewChanges": False,
            },
        },
    }


def _download_onlyoffice_update(download_url: str, target_path: Path) -> None:
    with urllib.request.urlopen(download_url) as response:
        data = response.read()
    target_path.write_bytes(data)


def _markdown_to_blocks(markdown: str) -> list[dict[str, object]]:
    blocks: list[dict[str, object]] = []
    lines = (markdown or "").splitlines()

    current_paragraph: list[str] = []
    current_list: dict[str, object] | None = None

    def flush_paragraph() -> None:
        nonlocal current_paragraph
        text = " ".join(part.strip() for part in current_paragraph if part.strip()).strip()
        if text:
            blocks.append({"type": "paragraph", "text": text})
        current_paragraph = []

    def flush_list() -> None:
        nonlocal current_list
        if current_list and current_list.get("items"):
            blocks.append(current_list)
        current_list = None

    for raw_line in lines:
        line = raw_line.rstrip()
        if not line.strip():
            flush_paragraph()
            flush_list()
            continue

        heading_match = re.match(r"^(#+)\s+(.*)", line)
        if heading_match:
            flush_paragraph()
            flush_list()
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            blocks.append({"type": "heading", "level": min(level, 5), "text": text})
            continue

        numbered_heading = re.match(r"^(\d+(?:\.\d+)*)\.\s+(.*)", line)
        if numbered_heading:
            flush_paragraph()
            flush_list()
            depth = numbered_heading.group(1).count(".") + 1
            text = numbered_heading.group(2).strip()
            blocks.append({"type": "heading", "level": min(depth, 5), "text": text})
            continue

        ordered_item = re.match(r"^\s*(\d+)[.)]\s+(.*)", line)
        if ordered_item:
            if current_list is None or current_list.get("type") != "ordered_list":
                flush_paragraph()
                flush_list()
                current_list = {"type": "ordered_list", "items": []}
            current_list["items"].append(ordered_item.group(2).strip())
            continue

        bullet_item = re.match(r"^\s*[-*+]\s+(.*)", line)
        if bullet_item:
            if current_list is None or current_list.get("type") != "bullet_list":
                flush_paragraph()
                flush_list()
                current_list = {"type": "bullet_list", "items": []}
            current_list["items"].append(bullet_item.group(1).strip())
            continue

        current_paragraph.append(line)

    flush_paragraph()
    flush_list()
    return blocks


def _blocks_to_html(blocks: list[dict[str, object]]) -> str:
    parts: list[str] = []

    for block in blocks:
        if block.get("type") == "heading":
            level = int(block.get("level", 1))
            text = html.escape(str(block.get("text", "")))
            level = max(1, min(level, 6))
            parts.append(f"<h{level}>{text}</h{level}>")
            continue

        if block.get("type") == "ordered_list":
            items = "".join(
                f"<li>{html.escape(str(item))}</li>" for item in block.get("items", [])
            )
            parts.append(f"<ol>{items}</ol>")
            continue

        if block.get("type") == "bullet_list":
            items = "".join(
                f"<li>{html.escape(str(item))}</li>" for item in block.get("items", [])
            )
            parts.append(f"<ul>{items}</ul>")
            continue

        text = html.escape(str(block.get("text", "")))
        if text:
            parts.append(f"<p>{text}</p>")

    return "\n".join(parts)


def _ensure_heading_numbering(document: Document) -> int:
    numbering_part = document.part.numbering_part
    numbering = numbering_part.numbering_definitions._numbering

    existing_ids = [
        int(num.get(qn("w:numId")))
        for num in numbering.findall(qn("w:num"))
        if num.get(qn("w:numId"))
    ]
    num_id = max(existing_ids or [0]) + 1
    abstract_num_id = num_id + 99

    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), str(abstract_num_id))

    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "hybridMultilevel")
    abstract_num.append(multi_level)

    for level in range(5):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))

        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)

        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "decimal")
        lvl.append(num_fmt)

        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), "%{}.".format(level + 1))
        lvl.append(lvl_text)

        p_style = OxmlElement("w:pStyle")
        p_style.set(qn("w:val"), f"Heading {level + 1}")
        lvl.append(p_style)

        indent = OxmlElement("w:ind")
        indent.set(qn("w:left"), str(360 * level))
        lvl.append(indent)

        abstract_num.append(lvl)

    numbering.append(abstract_num)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id_el = OxmlElement("w:abstractNumId")
    abstract_num_id_el.set(qn("w:val"), str(abstract_num_id))
    num.append(abstract_num_id_el)
    numbering.append(num)
    return num_id


def _attach_numbering(paragraph, num_id: int, level: int) -> None:
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(max(0, level)))
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(numid)
    paragraph._p.get_or_add_pPr().append(num_pr)


def _render_docx_from_blocks(blocks: list[dict[str, object]]) -> io.BytesIO:
    document = Document()
    numbering_id = _ensure_heading_numbering(document)

    for block in blocks or []:
        if block.get("type") == "heading":
            level = int(block.get("level", 1))
            style_level = max(1, min(level, 5))
            paragraph = document.add_paragraph(
                str(block.get("text", "")), style=f"Heading {style_level}"
            )
            _attach_numbering(paragraph, numbering_id, style_level - 1)
            continue

        if block.get("type") == "ordered_list":
            for item in block.get("items", []):
                document.add_paragraph(str(item), style="List Number")
            continue

        if block.get("type") == "bullet_list":
            for item in block.get("items", []):
                document.add_paragraph(str(item), style="List Bullet")
            continue

        document.add_paragraph(str(block.get("text", "")))

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def _find_block_node(node: NavigableString) -> Tag | None:
    current = node.parent
    while current is not None:
        if isinstance(current, Tag) and current.name in {
            "p",
            "li",
            "td",
            "th",
            "caption",
            "blockquote",
            "h1",
            "h2",
            "h3",
            "h4",
            "h5",
            "h6",
            "pre",
            "code",
            "div",
        }:
            return current
        current = current.parent
    return None


def _summarize_location(
    soup: BeautifulSoup,
    node_infos: list[dict[str, object]],
    start_index: int,
) -> DiffLocation | None:
    if not node_infos:
        return None

    target_info: dict[str, object] | None = None
    for info in node_infos:
        if info["start"] <= start_index < info["end"]:
            target_info = info
            break

    if target_info is None and start_index > 0:
        backtrack_index = start_index - 1
        for info in node_infos:
            if info["start"] <= backtrack_index < info["end"]:
                target_info = info
                break

    if target_info is None:
        for info in reversed(node_infos):
            if info["start"] <= start_index:
                target_info = info
                break

    if target_info is None:
        return None

    node = target_info["node"]
    if not isinstance(node, NavigableString):
        return None

    block = _find_block_node(node)
    block_summary: str | None = None
    if block is not None:
        text = block.get_text(" ", strip=True)
        if text:
            block_summary = _truncate_text(text)

    heading: Tag | None = None
    search_anchor: Tag | None = block if block is not None else node.parent
    heading_tags = ["h1", "h2", "h3", "h4", "h5", "h6"]
    if search_anchor is not None:
        heading = search_anchor if isinstance(search_anchor, Tag) and search_anchor.name in heading_tags else search_anchor.find_previous(heading_tags)
    else:
        heading = soup.find(heading_tags)

    section_title: str | None = None
    if heading is not None:
        section_text = heading.get_text(" ", strip=True)
        if section_text:
            section_title = _truncate_text(section_text, 60)

    if section_title is None and block_summary is None:
        return None

    return DiffLocation(section_title=section_title, block_summary=block_summary)


def _build_highlight_lookup(highlights: list[dict[str, object]]) -> dict[int, dict[str, object]]:
    lookup: dict[int, dict[str, object]] = {}
    for entry in highlights:
        start = entry["start"]
        end = entry["end"]
        for index in range(start, end):
            lookup[index] = entry
    return lookup


def _create_marker_tag(
    soup: BeautifulSoup, entry: dict[str, object], text: str
) -> Tag:
    mark = soup.new_tag("span")
    classes = [
        "diff-marker",
        f'diff-marker--{entry["type"]}',
        f'diff-marker--{entry["role"]}',
    ]
    if entry.get("placeholder"):
        classes.append("diff-marker--placeholder")
    else:
        classes.append("diff-marker--with-pill")
    mark["class"] = classes
    mark["data-diff-id"] = entry["id"]
    mark["data-diff-type"] = entry["type"]
    mark["data-diff-role"] = entry["role"]

    label = entry.get("label")
    number = entry.get("number")
    if label:
        mark["data-diff-type-label"] = label
    if number is not None:
        mark["data-diff-number"] = str(number)
    if label and number is not None:
        mark["title"] = f"{label} #{number}"
    if entry.get("placeholder"):
        mark["data-diff-placeholder"] = "true"

    mark.string = text
    return mark


def _apply_highlights(
    soup: BeautifulSoup,
    node_infos: list[dict[str, object]],
    highlights: list[dict[str, object]],
) -> str:
    if not highlights:
        return str(soup)

    span_highlights: list[dict[str, object]] = []
    boundary_highlights: dict[int, list[dict[str, object]]] = defaultdict(list)

    for entry in highlights:
        start = int(entry.get("start", 0))
        end = int(entry.get("end", start))
        if end > start:
            span_highlights.append(entry)
        else:
            boundary_highlights[start].append(entry)

    lookup = _build_highlight_lookup(span_highlights)

    for info in node_infos:
        node = info["node"]
        parent = node.parent
        if parent is None:
            continue

        tokens = info["tokens"]
        start_index = info["start"]
        end_index = info["end"]

        current_entry: dict[str, object] | None = None
        buffer: list[str] = []
        fragments: list[object] = []

        def emit_boundary(boundary_index: int) -> None:
            entries = boundary_highlights.pop(boundary_index, None)
            if not entries:
                return
            flush()
            for boundary_entry in entries:
                fragments.append(
                    _create_marker_tag(soup, boundary_entry, "\u00a0")
                )

        def flush() -> None:
            nonlocal buffer, current_entry
            if not buffer:
                return
            text = "".join(buffer)
            if current_entry:
                fragments.append(_create_marker_tag(soup, current_entry, text))
            else:
                fragments.append(text)
            buffer = []

        emit_boundary(start_index)
        for offset, token in enumerate(tokens):
            absolute_index = start_index + offset
            entry = lookup.get(absolute_index)
            if entry is not current_entry:
                flush()
                current_entry = entry
            buffer.append(token)
            emit_boundary(absolute_index + 1)

        flush()
        emit_boundary(end_index)

        for fragment in fragments:
            node.insert_before(fragment)

        node.extract()

    if boundary_highlights:
        fallback_parent: Tag | None = None
        if node_infos:
            fallback_parent = node_infos[-1]["node"].parent
        if fallback_parent is None:
            fallback_parent = soup.body or soup

        for boundary_index in sorted(boundary_highlights.keys()):
            entries = boundary_highlights[boundary_index]
            for entry in entries:
                fallback_parent.append(
                    _create_marker_tag(soup, entry, "\u00a0")
                )

    return str(soup)


def _build_diff(
    original_html: str, modified_html: str
) -> tuple[str, DiffStats, list[DiffItem], str, str]:
    from difflib import SequenceMatcher

    (
        original_soup,
        original_tokens,
        original_node_infos,
    ) = _prepare_html_tokens(original_html)
    (
        modified_soup,
        modified_tokens,
        modified_node_infos,
    ) = _prepare_html_tokens(modified_html)

    matcher = SequenceMatcher(None, original_tokens, modified_tokens)

    diff_parts: list[str] = []
    inserted_tokens = deleted_tokens = replaced_tokens = 0
    diff_items: list[DiffItem] = []
    highlight_map = {"original": [], "modified": []}
    diff_index = 1

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            diff_parts.append(_escape_tokens(original_tokens[i1:i2]))
        elif tag == "insert":
            if j1 == j2:
                continue
            diff_number = diff_index
            diff_id = f"diff-{diff_number}"
            diff_index += 1
            inserted_tokens += j2 - j1
            inserted_raw = "".join(modified_tokens[j1:j2])
            inserted_escaped = _escape_tokens(modified_tokens[j1:j2])
            modified_location = _summarize_location(
                modified_soup, modified_node_infos, j1
            )
            diff_parts.append(
                f'<ins class="diff-insert" data-diff-id="{diff_id}">{inserted_escaped}</ins>'
            )
            diff_items.append(
                DiffItem(
                    id=diff_id,
                    type="insert",
                    original_text="",
                    modified_text=inserted_raw,
                    modified_location=modified_location,
                )
            )
            highlight_map["original"].append(
                {
                    "id": diff_id,
                    "type": "insert",
                    "role": "original",
                    "start": i1,
                    "end": i1,
                    "label": DIFF_TYPE_LABELS["insert"],
                    "number": diff_number,
                    "placeholder": True,
                }
            )
            highlight_map["modified"].append(
                {
                    "id": diff_id,
                    "type": "insert",
                    "role": "modified",
                    "start": j1,
                    "end": j2,
                    "label": DIFF_TYPE_LABELS["insert"],
                    "number": diff_number,
                }
            )
        elif tag == "delete":
            if i1 == i2:
                continue
            diff_number = diff_index
            diff_id = f"diff-{diff_number}"
            diff_index += 1
            deleted_tokens += i2 - i1
            deleted_raw = "".join(original_tokens[i1:i2])
            deleted_escaped = _escape_tokens(original_tokens[i1:i2])
            original_location = _summarize_location(
                original_soup, original_node_infos, i1
            )
            # 在 modified 中找到对应的位置（删除后应该插入占位符的位置）
            # 由于是删除，modified 中对应的位置是 j1（等于 i1 在原始序列中的位置）
            # 但我们需要在 modified 的对应位置插入占位符
            modified_location = _summarize_location(
                modified_soup, modified_node_infos, j1 if j1 < len(modified_tokens) else max(0, len(modified_tokens) - 1)
            )
            diff_parts.append(
                f'<del class="diff-delete" data-diff-id="{diff_id}">{deleted_escaped}</del>'
            )
            diff_items.append(
                DiffItem(
                    id=diff_id,
                    type="delete",
                    original_text=deleted_raw,
                    modified_text="",
                    original_location=original_location,
                    modified_location=modified_location,
                )
            )
            highlight_map["original"].append(
                {
                    "id": diff_id,
                    "type": "delete",
                    "role": "original",
                    "start": i1,
                    "end": i2,
                    "label": DIFF_TYPE_LABELS["delete"],
                    "number": diff_number,
                }
            )
            # 在 modified 中添加占位符标记
            highlight_map["modified"].append(
                {
                    "id": diff_id,
                    "type": "delete",
                    "role": "modified",
                    "start": j1,
                    "end": j1,
                    "label": DIFF_TYPE_LABELS["delete"],
                    "number": diff_number,
                    "placeholder": True,
                }
            )
        elif tag == "replace":
            if i1 == i2 and j1 == j2:
                continue
            diff_number = diff_index
            diff_id = f"diff-{diff_number}"
            diff_index += 1
            removed_raw = "".join(original_tokens[i1:i2])
            added_raw = "".join(modified_tokens[j1:j2])
            removed_escaped = _escape_tokens(original_tokens[i1:i2])
            added_escaped = _escape_tokens(modified_tokens[j1:j2])
            original_location = _summarize_location(
                original_soup, original_node_infos, i1
            )
            modified_location = _summarize_location(
                modified_soup, modified_node_infos, j1
            )

            if i1 != i2:
                diff_parts.append(
                    f'<del class="diff-delete" data-diff-id="{diff_id}">{removed_escaped}</del>'
                )
                highlight_map["original"].append(
                    {
                        "id": diff_id,
                        "type": "replace",
                        "role": "original",
                        "start": i1,
                        "end": i2,
                        "label": DIFF_TYPE_LABELS["replace"],
                        "number": diff_number,
                    }
                )
            if j1 != j2:
                diff_parts.append(
                    f'<ins class="diff-insert" data-diff-id="{diff_id}">{added_escaped}</ins>'
                )
                highlight_map["modified"].append(
                    {
                        "id": diff_id,
                        "type": "replace",
                        "role": "modified",
                        "start": j1,
                        "end": j2,
                        "label": DIFF_TYPE_LABELS["replace"],
                        "number": diff_number,
                    }
                )

            replaced_tokens += max(i2 - i1, j2 - j1)
            diff_items.append(
                DiffItem(
                    id=diff_id,
                    type="replace",
                    original_text=removed_raw,
                    modified_text=added_raw,
                    original_location=original_location,
                    modified_location=modified_location,
                )
            )

    diff_html = "".join(diff_parts)
    stats = DiffStats(
        inserted_tokens=inserted_tokens,
        deleted_tokens=deleted_tokens,
        replaced_tokens=replaced_tokens,
    )
    highlighted_original = _apply_highlights(
        original_soup, original_node_infos, highlight_map["original"]
    )
    highlighted_modified = _apply_highlights(
        modified_soup, modified_node_infos, highlight_map["modified"]
    )

    return diff_html, stats, diff_items, highlighted_original, highlighted_modified


@app.post("/api/contract/desensitize", response_model=DesensitizeResponse)
async def desensitize_contract(
    file: Annotated[UploadFile, File(description="合同 .doc/.docx 文件")]
) -> DesensitizeResponse:
    raw_bytes, filename = await _read_word_file_bytes(file)
    plain_text = _collect_docx_text(raw_bytes)
    hits = _find_sensitive_hits(plain_text)

    mask_map = {hit.value: _mask_value(hit.value) for hit in hits if hit.value}
    sanitized_bytes = _sanitize_docx_bytes(raw_bytes, mask_map)
    sanitized_preview = _mask_text_with_map(plain_text, mask_map) if plain_text else None

    encoded_file = base64.b64encode(sanitized_bytes).decode("utf-8")
    safe_name = _sanitize_filename(filename.rsplit(".", 1)[0]) or "合同"

    return DesensitizeResponse(
        sanitized_docx=encoded_file,
        filename=f"{safe_name}_脱敏.docx",
        total_hits=sum(hit.count for hit in hits),
        hits=hits,
        sanitized_preview=sanitized_preview,
    )


@app.post("/convert", response_model=ConversionResponse)
async def convert_word(
    file: Annotated[UploadFile, File(description=".docx file to convert")]
) -> ConversionResponse:
    html, notes = await _convert_to_html(file)
    return ConversionResponse(html=html, notes=notes)


@app.post("/diff", response_model=DiffResponse)
async def diff_word_documents(
    original_file: Annotated[
        UploadFile,
        File(description="Original .docx file", alias="original"),
    ],
    modified_file: Annotated[
        UploadFile,
        File(description="Modified .docx file", alias="modified"),
    ],
) -> DiffResponse:
    original_html, original_notes = await _convert_to_html(original_file)
    modified_html, modified_notes = await _convert_to_html(modified_file)

    diff_html, stats, diff_items, highlighted_original, highlighted_modified = _build_diff(
        original_html, modified_html
    )

    return DiffResponse(
        original_html=highlighted_original or "<p>未检测到正文内容。</p>",
        modified_html=highlighted_modified or "<p>未检测到正文内容。</p>",
        diff_html=diff_html or "<p>未检测到差异。</p>",
        stats=stats,
        diff_items=diff_items,
        original_notes=original_notes,
        modified_notes=modified_notes,
    )


@app.post("/export")
async def export_document(request: ExportRequest) -> StreamingResponse:
    format_name = request.format.lower()
    filename = _sanitize_filename(request.filename)
    html_content = request.content or ""

    if format_name == "docx":
        buffer = _render_docx_document(html_content)
        media_type = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        extension = "docx"
    elif format_name == "pdf":
        buffer = _render_pdf_document(html_content)
        media_type = "application/pdf"
        extension = "pdf"
    elif format_name == "json":
        buffer = _render_json_document(html_content)
        media_type = "application/json"
        extension = "json"
    else:  # pragma: no cover - guarded by pydantic validation
        raise HTTPException(status_code=400, detail="Unsupported export format")

    download_name = f"{filename}.{extension}"
    data = buffer.getvalue()
    return _build_file_response(data, media_type, download_name)


@app.post("/api/ai/transform")
def ai_transform(request: AiTransformRequest) -> dict[str, str]:
    content = request.markdown.strip()
    if not content:
        return {"markdown": ""}

    if request.mode == "rewrite":
        return {"markdown": f"[改写示例]\n{content}"}

    if request.mode == "expand":
        return {
            "markdown": f"{content}\n\n扩写示例文本：增加履约节点、违约责任和沟通机制以确保条款可执行。"
        }

    if request.mode == "rephrase":
        return {"markdown": f"[重写示例] {content}"}

    safe_instruction = (request.user_instruction or "自定义指令").strip()
    return {
        "markdown": f"根据指令（{safe_instruction}）完成示例改写：\n{content}"
    }


@app.post("/api/export/docx")
def export_docx(payload: MarkdownPayload) -> StreamingResponse:
    blocks = _markdown_to_blocks(payload.markdown)
    html_content = _blocks_to_html(blocks)
    buffer = _render_docx_from_blocks(blocks)

    filename = _sanitize_filename("合同AI导出") + ".docx"
    data = buffer.getvalue()
    return _build_file_response(
        data,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename,
    )


@app.post("/onlyoffice/upload")
async def upload_onlyoffice(
    request: Request,
    file: Annotated[UploadFile, File(description="Office 文件")],
) -> dict[str, object]:
    file_id, display_name = await _persist_onlyoffice_upload(file)
    config = _build_onlyoffice_config(file_id, display_name, request)
    return {
        "fileId": file_id,
        "config": config,
        "documentServerUrl": DOCUMENT_SERVER_PUBLIC,
    }


@app.get("/onlyoffice/config/{file_id}")
async def onlyoffice_config(file_id: str, request: Request) -> dict[str, object]:
    file_path = _ensure_onlyoffice_file(file_id)
    display_name = file_path.name.split("_", 1)[1] if "_" in file_path.name else file_path.name
    config = _build_onlyoffice_config(file_id, display_name, request)
    return {
        "fileId": file_id,
        "config": config,
        "documentServerUrl": DOCUMENT_SERVER_PUBLIC,
    }


@app.get("/onlyoffice/files/{file_id}")
def onlyoffice_file(file_id: str) -> FileResponse:
    file_path = _ensure_onlyoffice_file(file_id)
    return FileResponse(file_path)


@app.post("/onlyoffice/callback/{file_id}")
async def onlyoffice_callback(file_id: str, request: Request) -> dict[str, int]:
    file_path = _ensure_onlyoffice_file(file_id)
    payload = await request.json()
    status = int(payload.get("status", 0))
    download_url = payload.get("url") or payload.get("changesurl")

    if status in {2, 3, 6} and download_url:
        try:
            _download_onlyoffice_update(download_url, file_path)
        except Exception:
            return {"error": 1}

    return {"error": 0}


@app.get("/health")
def health_check() -> dict[str, str]:
    return {"status": "ok"}
